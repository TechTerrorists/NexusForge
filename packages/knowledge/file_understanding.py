from __future__ import annotations

import os
import re
from pathlib import Path


class FileUnderstanding:
    """Extract and process text from various file formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json", ".xml", ".html", ".rtf"}

    def extract_text(self, file_path: str) -> str:
        """Extract text content from a file based on its extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}")

        if extension == ".pdf":
            return self._extract_pdf(path)
        elif extension in (".docx", ".doc"):
            return self._extract_docx(path)
        elif extension in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".rtf"):
            return self._extract_text_file(path)
        return ""

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            import pymupdf
            doc = pymupdf.open(str(path))
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(path))
                text_parts = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_parts.append(extracted)
                return "\n".join(text_parts)
            except ImportError:
                raise ImportError(
                    "PDF extraction requires pymupdf or PyPDF2. "
                    "Install with: pip install pymupdf or pip install PyPDF2"
                )

    def _extract_docx(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            doc = Document(str(path))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("DOCX extraction requires python-docx. Install with: pip install python-docx")

    def _extract_text_file(self, path: Path) -> str:
        """Extract text from plain text and markup files."""
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="latin-1", errors="replace") as f:
                return f.read()

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 512,
        overlap: int = 50,
    ) -> list[str]:
        """Split text into overlapping chunks by character count."""
        if not text or not text.strip():
            return []
        text = text.strip()
        if len(text) <= chunk_size:
            return [text]
        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk.strip())
            start = end - overlap
            if start >= len(text):
                break
        return chunks

    def extract_citations(self, text: str) -> list[dict[str, str]]:
        """Extract citation references from text using common patterns."""
        citations: list[dict[str, str]] = []
        seen: set[str] = set()

        bracket_pattern = re.compile(r"\[(\d+(?:[,;\s]+\d+)*)\]")
        for match in bracket_pattern.finditer(text):
            numbers = re.split(r"[,;\s]+", match.group(1))
            for num in numbers:
                num = num.strip()
                if num and num not in seen:
                    seen.add(num)
                    citations.append({"type": "bracket", "reference": num, "raw": match.group(0)})

        author_year_pattern = re.compile(
            r"\(([A-Z][a-z]+(?:\s+(?:et\s+al\.?|and\s+[A-Z][a-z]+))?),\s*(\d{4}[a-z]?)\)"
        )
        for match in author_year_pattern.finditer(text):
            author = match.group(1)
            year = match.group(2)
            key = f"{author}_{year}"
            if key not in seen:
                seen.add(key)
                citations.append({
                    "type": "author_year",
                    "author": author,
                    "year": year,
                    "raw": match.group(0),
                })

        doi_pattern = re.compile(r"\b(10\.\d{4,9}/[^\s]+)")
        for match in doi_pattern.finditer(text):
            doi = match.group(1)
            if doi not in seen:
                seen.add(doi)
                citations.append({"type": "doi", "reference": doi, "raw": match.group(0)})

        url_pattern = re.compile(r"https?://[^\s<>\"')\]]+")
        for match in url_pattern.finditer(text):
            url = match.group(0)
            if url not in seen:
                seen.add(url)
                citations.append({"type": "url", "reference": url, "raw": url})

        return citations
